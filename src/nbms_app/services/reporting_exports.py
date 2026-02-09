from __future__ import annotations

import hashlib
import json
from io import BytesIO

from django.core.exceptions import ValidationError
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.template.loader import render_to_string
from django.utils import timezone

from nbms_app.models import (
    AnnexSectionResponse,
    ReportExportArtifact,
    ReportTemplatePackResponse,
)
from nbms_app.services.reporting_workflow import resolve_cbd_pack


def _canonical_json(value):
    return json.dumps(value or {}, sort_keys=True, ensure_ascii=True, separators=(",", ":"))


def _hash_bytes(value):
    return hashlib.sha256(value).hexdigest()


def _hash_payload(value):
    return hashlib.sha256(_canonical_json(value).encode("utf-8")).hexdigest()


def _normalise_response_rows(instance, pack):
    section_rows = list(pack.sections.filter(is_active=True).order_by("ordering", "code"))
    response_map = {
        row.section_id: row.response_json
        for row in ReportTemplatePackResponse.objects.filter(
            reporting_instance=instance,
            section__pack=pack,
        ).select_related("section")
    }
    sections = []
    for section in section_rows:
        sections.append(
            {
                "code": section.code,
                "title": section.title,
                "ordering": section.ordering,
                "content": response_map.get(section.id, {}),
            }
        )
    return sections


def build_cbd_report_payload(*, instance):
    pack = resolve_cbd_pack()
    sections = _normalise_response_rows(instance, pack)
    annex_rows = (
        AnnexSectionResponse.objects.filter(reporting_instance=instance, is_enabled=True)
        .order_by("ordering", "decision_topic_code", "id")
    )
    annex = [
        {
            "decision_topic_code": row.decision_topic_code,
            "title": row.title,
            "ordering": row.ordering,
            "response_json": row.response_json,
        }
        for row in annex_rows
    ]
    payload = {
        "schema": "nbms.cbd_national_report.v1",
        "exporter_version": "1.0.0",
        "generated_at": timezone.now().isoformat(),
        "reporting_instance": {
            "uuid": str(instance.uuid),
            "cycle_code": instance.cycle.code if instance.cycle_id else "",
            "cycle_title": instance.cycle.title if instance.cycle_id else "",
            "version_label": instance.version_label,
            "report_title": instance.report_title,
            "country_name": instance.country_name,
            "is_public": bool(instance.is_public),
            "status": instance.status,
            "final_content_hash": instance.final_content_hash,
            "focal_point_org": instance.focal_point_org.name if instance.focal_point_org_id else "",
            "publishing_authority_org": instance.publishing_authority_org.name if instance.publishing_authority_org_id else "",
        },
        "sections": sections,
        "annex": annex,
    }
    payload["payload_hash"] = _hash_payload(payload)
    return payload


def render_cbd_pdf_bytes(*, payload):
    try:
        from xhtml2pdf import pisa  # noqa: WPS433
    except Exception as exc:  # noqa: BLE001
        raise ValidationError(f"PDF renderer dependency missing: {exc}") from exc

    context = {
        "payload": payload,
        "instance": payload.get("reporting_instance", {}),
        "sections": payload.get("sections", []),
        "annex": payload.get("annex", []),
        "generated_at": timezone.now(),
    }
    html = render_to_string("nbms_app/reporting/national_report_pdf.html", context)
    output = BytesIO()
    result = pisa.CreatePDF(src=html, dest=output, encoding="utf-8")
    if result.err:
        raise ValidationError("Failed to render national report PDF.")
    return output.getvalue()


def render_cbd_docx_bytes(*, payload):
    try:
        from docx import Document  # noqa: WPS433
    except Exception as exc:  # noqa: BLE001
        raise ValidationError(f"DOCX renderer dependency missing: {exc}") from exc

    doc = Document()
    report = payload.get("reporting_instance", {})
    doc.add_heading(report.get("report_title") or "CBD National Report", level=0)
    meta_line = f"{report.get('country_name', '')} | {report.get('cycle_code', '')} | {report.get('version_label', '')}"
    doc.add_paragraph(meta_line)
    doc.add_paragraph(f"Visibility: {'Public' if report.get('is_public') else 'Internal'}")

    for section in payload.get("sections", []):
        doc.add_heading(f"{section.get('code', '').upper()} - {section.get('title', '')}", level=1)
        content = section.get("content") or {}
        for key in sorted(content.keys()):
            value = content.get(key)
            doc.add_paragraph(f"{key}:", style="List Bullet")
            if isinstance(value, list):
                if value and isinstance(value[0], dict):
                    for row in value:
                        doc.add_paragraph(json.dumps(row, ensure_ascii=False), style="List Number")
                else:
                    doc.add_paragraph(", ".join(str(item) for item in value))
            elif isinstance(value, dict):
                for row_key in sorted(value.keys()):
                    doc.add_paragraph(f"{row_key}: {value.get(row_key)}")
            else:
                doc.add_paragraph(str(value or ""))

    if payload.get("annex"):
        doc.add_heading("Annex", level=1)
        for annex in payload.get("annex", []):
            doc.add_heading(f"{annex.get('decision_topic_code')} - {annex.get('title')}", level=2)
            response_json = annex.get("response_json") or {}
            for key in sorted(response_json.keys()):
                doc.add_paragraph(f"{key}: {response_json.get(key)}")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def store_report_export_artifact(*, instance, generated_by, format_name, content_bytes, linked_action=None, metadata=None):
    timestamp = timezone.now().strftime("%Y%m%dT%H%M%S")
    suffix = "bin"
    if format_name == ReportExportArtifact.FORMAT_PDF:
        suffix = "pdf"
    elif format_name == ReportExportArtifact.FORMAT_DOCX:
        suffix = "docx"
    elif format_name == ReportExportArtifact.FORMAT_JSON:
        suffix = "json"
    elif format_name == ReportExportArtifact.FORMAT_DOSSIER:
        suffix = "zip"
    storage_path = f"reports/{instance.uuid}/{timestamp}_{format_name}.{suffix}"
    if default_storage.exists(storage_path):
        default_storage.delete(storage_path)
    default_storage.save(storage_path, ContentFile(content_bytes))
    digest = _hash_bytes(content_bytes)
    artifact = ReportExportArtifact.objects.create(
        reporting_instance=instance,
        format=format_name,
        storage_path=storage_path,
        content_hash=digest,
        generated_by=generated_by if getattr(generated_by, "is_authenticated", False) else None,
        immutable=bool(instance.finalized_at),
        linked_action=linked_action,
        metadata_json=metadata or {},
    )
    return artifact

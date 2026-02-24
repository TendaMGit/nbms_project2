from __future__ import annotations

import hashlib
import json
import re
from io import BytesIO
from urllib import request as urllib_request

from django.core.exceptions import ValidationError
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.utils import timezone

from nbms_app.models import (
    Indicator,
    IndicatorDataPoint,
    MonitoringProgramme,
    ReportContext,
    ReportNarrativeBlock,
    ReportNarrativeBlockVersion,
    ReportTemplatePack,
    ReportTemplatePackResponse,
    ReportingInstance,
    TaxonGoldSummary,
    EcosystemGoldSummary,
    IASGoldSummary,
)


TOKEN_PATTERN = re.compile(r"\[\[nbms:(?P<token_type>[a-zA-Z0-9_]+)\s*(?P<args>[^\]]*)\]\]")


def _canonical_json(value):
    return json.dumps(value or {}, sort_keys=True, separators=(",", ":"), ensure_ascii=True)


def _hash_payload(value):
    return hashlib.sha256(_canonical_json(value).encode("utf-8")).hexdigest()


def _hash_bytes(value):
    return hashlib.sha256(value).hexdigest()


def _safe_text(value):
    if value is None:
        return ""
    return str(value).strip()


def _parse_args(args):
    parsed = {}
    for part in (args or "").split():
        if "=" not in part:
            continue
        key, value = part.split("=", 1)
        parsed[key.strip().lower()] = value.strip().strip('"').strip("'")
    return parsed


def _resolve_pack():
    for code in ("cbd_national_report_v1", "cbd_ort_nr7_v2"):
        pack = ReportTemplatePack.objects.filter(code=code, is_active=True).first()
        if pack:
            return pack
    return None


def normalize_context_filters(raw):
    raw = raw or {}
    if isinstance(raw, str):
        try:
            raw = json.loads(raw)
        except Exception:  # noqa: BLE001
            raw = {}
    if not isinstance(raw, dict):
        raw = {}
    return {
        "report_label": _safe_text(raw.get("report_label")),
        "period_start": _safe_text(raw.get("period_start")),
        "period_end": _safe_text(raw.get("period_end")),
        "geography": _safe_text(raw.get("geography")),
        "disaggregation": _safe_text(raw.get("disaggregation")),
        "indicator_code": _safe_text(raw.get("indicator_code")),
        "programme_code": _safe_text(raw.get("programme_code")),
    }


def persist_report_context(*, instance, user=None, session_key="", filters=None):
    filters = normalize_context_filters(filters)
    context, _created = ReportContext.objects.update_or_create(
        reporting_instance=instance,
        user=user if getattr(user, "is_authenticated", False) else None,
        session_key=session_key or "",
        defaults={
            "filters_json": filters,
            "context_hash": _hash_payload(filters),
            "is_active": True,
            "updated_by": user if getattr(user, "is_authenticated", False) else None,
        },
    )
    return context


def _indicator_value(code, year=None):
    indicator = Indicator.objects.filter(code=code).first()
    if not indicator:
        return None
    queryset = IndicatorDataPoint.objects.filter(series__indicator=indicator)
    if year:
        try:
            queryset = queryset.filter(year=int(year))
        except Exception:  # noqa: BLE001
            pass
    point = queryset.order_by("-year", "-id").first()
    if not point:
        return None
    if point.value_numeric is not None:
        return float(point.value_numeric)
    return point.value_text


def _registry_value(kind, metric):
    model = {
        "taxa": TaxonGoldSummary,
        "ecosystems": EcosystemGoldSummary,
        "ias": IASGoldSummary,
    }.get((kind or "").strip().lower())
    if not model:
        return None
    row = model.objects.order_by("-snapshot_date", "-id").first()
    if not row:
        return None
    return getattr(row, metric, None)


def _programme_value(code, metric):
    programme = MonitoringProgramme.objects.filter(programme_code=code).first()
    if not programme:
        return None
    if metric == "last_run_at":
        return programme.last_run_at.isoformat() if programme.last_run_at else ""
    if metric == "title":
        return programme.title
    return getattr(programme, metric, None)


def _resolve_token(token_type, args, context):
    if token_type == "indicator":
        indicator_code = args.get("code") or context.get("indicator_code")
        year = args.get("year")
        value = _indicator_value(indicator_code, year=year)
        return {
            "value": value if value is not None else "",
            "source": "indicator_data_point",
            "params": {"code": indicator_code, "year": year},
        }
    if token_type == "registry":
        kind = args.get("kind")
        metric = args.get("metric") or "profile_count"
        value = _registry_value(kind, metric)
        return {
            "value": value if value is not None else "",
            "source": f"registry_gold:{kind}",
            "params": {"kind": kind, "metric": metric},
        }
    if token_type == "programme":
        programme_code = args.get("code") or context.get("programme_code")
        metric = args.get("metric") or "title"
        value = _programme_value(programme_code, metric)
        return {
            "value": value if value is not None else "",
            "source": "monitoring_programme",
            "params": {"code": programme_code, "metric": metric},
        }
    return {"value": "", "source": "unsupported", "params": args}


def resolve_narrative_tokens(*, html, context_filters=None):
    context = normalize_context_filters(context_filters)
    manifest = []

    def _replacement(match):
        token_type = match.group("token_type").strip().lower()
        args = _parse_args(match.group("args"))
        resolved = _resolve_token(token_type, args, context)
        token_text = match.group(0)
        value_text = _safe_text(resolved.get("value"))
        manifest.append(
            {
                "token": token_text,
                "token_type": token_type,
                "params": resolved.get("params", {}),
                "resolved_value": value_text,
                "source": resolved.get("source"),
                "context_hash": _hash_payload(context),
            }
        )
        return value_text

    rendered = TOKEN_PATTERN.sub(_replacement, html or "")
    return {
        "rendered_html": rendered,
        "resolved_values_manifest": manifest,
        "context": context,
        "context_hash": _hash_payload(context),
    }


def _extract_docx_text(content_bytes):
    try:
        from docx import Document  # noqa: WPS433
    except Exception:  # noqa: BLE001
        return ""
    document = Document(BytesIO(content_bytes))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n".join(lines)


def _text_to_html(text):
    if not text:
        return ""
    chunks = [line.strip() for line in text.splitlines() if line.strip()]
    return "".join([f"<p>{line}</p>" for line in chunks])


def _build_initial_docx_bytes(title, text=""):
    try:
        from docx import Document  # noqa: WPS433
    except Exception as exc:  # noqa: BLE001
        raise ValidationError(f"DOCX dependency missing: {exc}") from exc
    document = Document()
    document.add_heading(title, level=1)
    if text:
        document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_docx_bytes_from_text(*, title, text=""):
    return _build_initial_docx_bytes(title=title, text=text)


def ensure_narrative_block(*, instance, section_code, block_key="main", title="Narrative", user=None):
    block, _created = ReportNarrativeBlock.objects.get_or_create(
        reporting_instance=instance,
        section_code=section_code,
        block_key=block_key,
        defaults={"title": title},
    )
    if block.current_version > 0 and block.storage_path:
        return block
    initial_bytes = _build_initial_docx_bytes(title=title, text="")
    upsert_narrative_block_content(
        block=block,
        content_bytes=initial_bytes,
        user=user,
        note="initial",
    )
    return block


def upsert_narrative_block_content(*, block, content_bytes, user=None, note=""):
    now = timezone.now()
    next_version = int(block.current_version or 0) + 1
    digest = _hash_bytes(content_bytes)
    storage_path = (
        f"reports/{block.reporting_instance.uuid}/narratives/"
        f"{block.section_code}_{block.block_key}_v{next_version}.docx"
    )
    if default_storage.exists(storage_path):
        default_storage.delete(storage_path)
    default_storage.save(storage_path, ContentFile(content_bytes))

    text_snapshot = _extract_docx_text(content_bytes)
    html_snapshot = _text_to_html(text_snapshot)

    version_row = ReportNarrativeBlockVersion.objects.create(
        block=block,
        version=next_version,
        storage_path=storage_path,
        content_hash=digest,
        html_snapshot=html_snapshot,
        text_snapshot=text_snapshot,
        created_by=user if getattr(user, "is_authenticated", False) else None,
    )
    block.storage_path = storage_path
    block.current_version = next_version
    block.current_content_hash = digest
    block.html_snapshot = html_snapshot
    block.text_snapshot = text_snapshot
    block.onlyoffice_document_key = f"{block.uuid}-{next_version}"
    block.updated_by = user if getattr(user, "is_authenticated", False) else None
    block.save(
        update_fields=[
            "storage_path",
            "current_version",
            "current_content_hash",
            "html_snapshot",
            "text_snapshot",
            "onlyoffice_document_key",
            "updated_by",
            "updated_at",
        ]
    )
    return version_row


def update_narrative_block_from_callback(*, block, callback_payload, user=None):
    callback_payload = callback_payload or {}
    download_url = callback_payload.get("url")
    if not download_url:
        return None
    with urllib_request.urlopen(download_url, timeout=15) as response:  # noqa: S310
        content_bytes = response.read()
    return upsert_narrative_block_content(
        block=block,
        content_bytes=content_bytes,
        user=user,
        note="onlyoffice_callback",
    )


def list_section_narrative_blocks(*, instance, section_code):
    return ReportNarrativeBlock.objects.filter(
        reporting_instance=instance,
        section_code=section_code,
    ).order_by("block_key")


def section_narrative_html(*, instance, section_code):
    blocks = list_section_narrative_blocks(instance=instance, section_code=section_code)
    if not blocks:
        return ""
    return "".join([row.html_snapshot for row in blocks if row.html_snapshot])


def build_section_chart_specs(*, instance, section_code, context_filters=None):
    context = normalize_context_filters(context_filters)
    pack = _resolve_pack()
    if not pack:
        return {"charts": [], "context": context, "context_hash": _hash_payload(context)}
    section = pack.sections.filter(code=section_code, is_active=True).first()
    if not section:
        return {"charts": [], "context": context, "context_hash": _hash_payload(context)}
    response = ReportTemplatePackResponse.objects.filter(
        reporting_instance=instance,
        section=section,
    ).first()
    payload = response.response_json or {} if response else {}
    charts = []

    if section_code == "section-iii":
        rows = payload.get("target_progress_rows") or []
        buckets = {}
        for row in rows:
            key = _safe_text(row.get("progress_level") or "unknown") or "unknown"
            buckets[key] = buckets.get(key, 0) + 1
        labels = sorted(buckets.keys())
        charts.append(
            {
                "id": "section3-progress-levels",
                "title": "Section III Progress Levels",
                "spec": {
                    "data": [{"type": "bar", "x": labels, "y": [buckets[label] for label in labels], "marker": {"color": "#2c7a5a"}}],
                    "layout": {"margin": {"l": 40, "r": 20, "t": 40, "b": 40}, "height": 300},
                    "config": {"displayModeBar": False, "responsive": True},
                },
            }
        )

    elif section_code == "section-iv":
        rows = payload.get("target_progress_rows") or []
        labels = []
        values = []
        for row in rows[:12]:
            labels.append(_safe_text(row.get("framework_target_code") or row.get("target_code") or "target"))
            values.append(1 if _safe_text(row.get("progress_level")).lower() in {"on_track", "achieved"} else 0)
        charts.append(
            {
                "id": "section4-target-status",
                "title": "Section IV Target Status (On-track=1)",
                "spec": {
                    "data": [{"type": "scatter", "mode": "lines+markers", "x": labels, "y": values, "line": {"color": "#15506c"}}],
                    "layout": {"margin": {"l": 40, "r": 20, "t": 40, "b": 60}, "height": 320},
                    "config": {"displayModeBar": False, "responsive": True},
                },
            }
        )

    elif section_code in {"section-i", "section-ii", "section-v"}:
        filled = 0
        total = 0
        for key, value in payload.items():
            if key.startswith("_"):
                continue
            total += 1
            if isinstance(value, list):
                filled += 1 if value else 0
            elif isinstance(value, dict):
                filled += 1 if value else 0
            elif _safe_text(value):
                filled += 1
        charts.append(
            {
                "id": f"{section_code}-completion",
                "title": "Structured Field Completion",
                "spec": {
                    "data": [
                        {
                            "type": "pie",
                            "labels": ["Completed", "Missing"],
                            "values": [filled, max(total - filled, 0)],
                            "marker": {"colors": ["#1f7a4d", "#d8e8dd"]},
                            "textinfo": "label+percent",
                        }
                    ],
                    "layout": {"margin": {"l": 20, "r": 20, "t": 40, "b": 20}, "height": 280},
                    "config": {"displayModeBar": False, "responsive": True},
                },
            }
        )

    return {"charts": charts, "context": context, "context_hash": _hash_payload(context)}


def render_section_narrative(*, instance, section_code, context_filters=None):
    raw_html = section_narrative_html(instance=instance, section_code=section_code)
    resolved = resolve_narrative_tokens(html=raw_html, context_filters=context_filters)
    return {
        "section_code": section_code,
        "raw_html": raw_html,
        "rendered_html": resolved["rendered_html"],
        "resolved_values_manifest": resolved["resolved_values_manifest"],
        "context": resolved["context"],
        "context_hash": resolved["context_hash"],
    }
